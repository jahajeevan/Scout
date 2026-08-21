"""Text extraction from uploaded files (spec §14).

Turns raw bytes into a list of ``Segment``s — text plus a page number where the
format has one (PDF) so citations can say "page 12". Supported: PDF (pypdf), DOCX
(python-docx), and any UTF-8 text/data/code file (txt, md, csv, json, xml, source).
Unsupported or unreadable files raise ``ExtractError`` with a human message.
"""

from __future__ import annotations

from dataclasses import dataclass


class ExtractError(Exception):
    """Raised when a file can't be read as text (bad type, corrupt, empty)."""


@dataclass
class Segment:
    text: str
    page: int | None = None  # 1-based page for paged formats, else None


# Extensions we treat as plain UTF-8 text (decode-and-go).
_TEXT_EXTS = {
    "txt", "md", "markdown", "csv", "tsv", "json", "xml", "yaml", "yml", "log", "rtf",
    "py", "js", "ts", "tsx", "jsx", "java", "c", "cpp", "h", "hpp", "cs", "go", "rs",
    "rb", "php", "swift", "kt", "sh", "sql", "html", "css", "toml", "ini", "env",
}


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def extract(filename: str, data: bytes) -> list[Segment]:
    """Extract text segments from a file's bytes. Raises ExtractError on failure."""
    ext = _ext(filename)
    if ext == "pdf":
        return _extract_pdf(data)
    if ext in ("docx", "dotx"):
        return _extract_docx(data)
    if ext in _TEXT_EXTS or not ext:
        return _extract_text(data)
    # Last resort: try to decode as text; if it's binary, fail honestly.
    try:
        return _extract_text(data)
    except ExtractError:
        raise ExtractError(f"Unsupported file type: .{ext}")


def _extract_text(data: bytes) -> list[Segment]:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ExtractError("File isn't readable as text.")
    text = text.strip()
    if not text:
        raise ExtractError("File is empty.")
    return [Segment(text=text, page=None)]


def _extract_pdf(data: bytes) -> list[Segment]:
    try:
        import io

        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractError(f"Couldn't open the PDF ({str(exc)[:80]}).")
    segments: list[Segment] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if text:
            segments.append(Segment(text=text, page=i))
    if not segments:
        raise ExtractError("No extractable text in the PDF (it may be scanned images).")
    return segments


def _extract_docx(data: bytes) -> list[Segment]:
    try:
        import io

        import docx

        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractError(f"Couldn't open the Word document ({str(exc)[:80]}).")
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include table cell text too — often where the real content lives.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractError("The Word document has no readable text.")
    return [Segment(text=text, page=None)]
